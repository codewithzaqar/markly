from PyQt5.QtWidgets import QFileDialog, QMessageBox
from markdown2 import markdown
from PyQt5.QtPrintSupport import QPrinter
from PyQt5.QtGui import QTextDocument
import docx

class ExportManager:
    def __init__(self, parent):
        self.parent = parent

    def export_to_html(self):
        """Export the Markdown content to an HTML file."""
        file_path, _ = QFileDialog.getSaveFileName(self.parent, "Export to HTML", "", "HTML Files (*.html);;All Files (*)")
        if file_path:
            try:
                markdown_text = self.parent.editor.toPlainText()
                html_content = markdown(markdown_text, extras=["tables", "footnotes"])
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(html_content)
                QMessageBox.information(self.parent, "Success", "File exported to HTML successfully!")
            except Exception as e:
                QMessageBox.critical(self.parent, "Error", f"Failed to export to HTML: {e}")

    def export_to_pdf(self):
        """Export the Markdown content to a PDF file."""
        file_path, _ = QFileDialog.getSaveFileName(self.parent, "Export to PDF", "", "PDF Files (*.pdf);;All Files (*)")
        if file_path:
            try:
                markdown_text = self.parent.editor.toPlainText()
                html_content = markdown(markdown_text, extras=["tables", "footnotes"])

                # Create a QTextDocument to render the HTML
                document = QTextDocument()
                document.setHtml(html_content)

                # Set up the printer
                printer = QPrinter()
                printer.setOutputFormat(QPrinter.PdfFormat)
                printer.setOutputFileName(file_path)

                # Print the document to the PDF file
                document.print_(printer)
                QMessageBox.information(self.parent, "Success", "File exported to PDF successfully!")
            except Exception as e:
                QMessageBox.critical(self.parent, "Error", f"Failed to export to PDF: {e}")

    def export_to_word(self):
        """Export the Markdown content to a Word file."""
        file_path, _ = QFileDialog.getSaveFileName(self.parent, "Export to Word", "", "Word Files (*.docx);;All Files (*)")
        if file_path:
            try:
                markdown_text = self.parent.editor.toPlainText()
                html_content = markdown(markdown_text, extras=["tables", "footnotes"])

                # Create a new Word document
                doc = docx.Document()
                doc.add_paragraph(html_content)
                doc.save(file_path)
                QMessageBox.information(self.parent, "Success", "File exported to Word successfully!")
            except Exception as e:
                QMessageBox.critical(self.parent, "Error", f"Failed to export to Word: {e}")